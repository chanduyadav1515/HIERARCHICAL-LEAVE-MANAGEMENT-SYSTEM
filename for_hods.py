from docx import Document
from tkinter import *
import tkinter
import os
from tkinter import messagebox
from openpyxl import load_workbook
import requests
global c
fac = tkinter.Tk()
fac.title("HEAD OF THE DEPARTMENT")
fac.geometry('1366x768')
fac.configure(bg='WHITE')
fac.iconbitmap('snake.ico')

file_url = "http://tejusolutions.com/rohit/download.php?filename=index.xlsx"
  
r = requests.get(file_url, stream = True) 
  
with open("index.xlsx","wb") as xlsx: 
    for chunk in r.iter_content(chunk_size=1024): 
        if chunk: 
            xlsx.write(chunk)


wb = load_workbook('index.xlsx')
# functions

def next_file():
    global c,row_scan,r_prev
    column = department()   
    lnth = len(sheet['A'])
    row_scan = r_prev
    while((row_scan+1)<=lnth):
        column = department()
        mat = column+str(row_scan+1)
        Value = sheet[mat].value
        if (Value == 1):
        # Get the file ID
            column = "A"
            f_id = column+str(row_scan+1)
            f_name = sheet[f_id].value
            b  = ".docx"
            c=f_name+b
            lBody_text.configure(text=" ")
            my_function(c)
            r_prev = row_scan+1
            break
        row_scan+=1
    else:
        messagebox.showinfo("DONE","No Files Pending")
        fac.destroy()
 
def accept():
    messagebox.showinfo("Forwarded", "Your response has been forwarded to PRINCIPAL.\n                                 Thankyou")
    coment = tCom.get("1.0","end-1c")
    global document
    document.add_paragraph(str(coment))
    global c
    document.save(c) # save comments to the word file  - if any
    column = department()
    global row_scan
    cell_no = column+str(row_scan)
    sheet[cell_no] = 2
    if (sheet["G"+str(row_scan)].value ==0):
        sheet["G"+str(row_scan)] =1
    elif (sheet["H"+str(row_scan)].value ==0):
        sheet["H"+str(row_scan)] =1
    elif (sheet["I"+str(row_scan)].value ==0):
        sheet["I"+str(row_scan)] =1
    elif (sheet["J"+str(row_scan)].value ==0):
        sheet["J"+str(row_scan)] =1
    wb.save('index.xlsx')
    lnth = len(sheet['A'])
    global r_prev
    row_scan = r_prev
    while(row_scan<=lnth):
        column = department()
        mat = column+str(row_scan)
        Value = sheet[mat].value
        if (Value == 1):
        # Get the file ID
            column = "A"
            f_id = column+str(row_scan)
            f_name = sheet[f_id].value
            b  = ".docx"
            c=f_name+b
            lBody_text.configure(text=" ")
            my_function(c)
            r_prev = row_scan
            break         
        row_scan+=1
    else:
        messagebox.showinfo("DONE","No Files Pending")
        fac.destroy()
def reject():
    messagebox.showerror("Rejected", "Your response has been submitted successfully.\n                                 Thankyou")
    coment = tCom.get("1.0","end-1c")
    global document
    document.add_paragraph(str(coment))
    global c
    document.save(c) # save comments to the word file  - if any
    column = department()
    global row_scan
    cell_no = column+str(row_scan)  # for HOD
    sheet[cell_no] = 3
    cell_no = "B"+str(row_scan)     # for faculty - 
    sheet[cell_no] = 8   
    wb.save('index.xlsx')
    lnth = len(sheet['A'])
    global r_prev
    row_scan = r_prev
    while(row_scan<=lnth):
        column = department()
        mat = column+str(row_scan)
        Value = sheet[mat].value
        if (Value == 1):
        # Get the file ID
            column = "A"
            f_id = column+str(row_scan)
            f_name = sheet[f_id].value
            b  = ".docx"
            c=f_name+b
            my_function(c)
            r_prev = row_scan
            break        
        row_scan+=1
    else:
        messagebox.showinfo("DONE3","No Files Pending")
        fac.destroy()

def open_file():

    os.startfile(c)

def start():
    lnth = len(sheet['A'])
    global row_scan
    row_scan = 2
    while(row_scan<=lnth):
        column = department()
        mat = column+str(row_scan)
        Value = sheet[mat].value
        if (Value == 1):
            # Get the file ID
            column = "A"
            f_id = column+str(row_scan)
            f_name = sheet[f_id].value
            b = ".docx"
            global c
            c=f_name+b
            download_file(c)
            my_function(c)
            global r_prev
            r_prev = row_scan
            lsel.place_forget()
            ld.place_forget()
            bs.place_forget()
            break
        row_scan += 1
    else:
        messagebox.showinfo("DONE1","No Files Pending")
        fac.destroy()

def download_file(c):
            file_url = "http://tejusolutions.com/rohit/download.php"
            url = file_url+'?filename='+c
            r = requests.get(url, stream=True)
            with open(c, "wb") as download_file:    
                for chunk in r.iter_content(chunk_size=1024):
                    if chunk:
                        download_file.write(chunk)


def my_function(c):
        global document
        print(c)
        document = Document(c) # open that respective word Document
        #get values from word file
        sub = document.paragraphs[2].text
        body = document.paragraphs[4].text
        date = document.paragraphs[0].text
        From = document.paragraphs[6].text
        lFrom=Label(fac, text ="SUBMITTED BY :   ", background = "white", fg="blue",font = "Castellar 12 ")
        lFrom.place(x=50, y=20)
        lFrom_text=Label(fac,text=From, background = "white", fg="black", justify="left",font = "times 12")
        lFrom_text.place(x=240, y=20)
        lDate=Label(fac, text ="Date Submitted : ", background = "white", fg="blue",font = "Helvetica 12")
        lDate.place(x=800, y=20)
        lDate_text=Label(fac,text=date, background = "white", fg="black", justify="left",font = "times 12")
        lDate_text.place(x=950, y=22)
            
        lSub=Label(fac, text ="SUBJECT            : ", background = "white", fg="firebrick3",font = "Algerian 14 bold ")
        lSub.place(x=50, y=100)
        lSubject=Label(fac,text=sub, wraplength=1000, background = "white", fg="OrangeRed3", justify="left",font = "times 14")
        lSubject.place(x=270, y=100)
        global lBody_text
        lBody=Label(fac, text ="CONTENT", background = "white", fg="DarkOrchid2",font = "Helvetica 14 bold")
        lBody.place(x=50, y=170)
        lBody_text=Label(fac,text=' ', wraplength=1000, background = "white", fg="black", justify="left",font = "times 13")
        lBody_text.configure(text=body)
        lBody_text.place(x=120, y=220)
      
        lCom=Label(fac, text ="COMMENTS", background = "white", fg="blue",font = "Helvetica 14 bold")
        lCom.place(x=50, y=400)
       
        ba = Button(fac, image = p_a, command=accept, bg = "white", highlightthickness = 0, bd = 0)
        ba.place(x = 350, y = 550)
      
        bReject = Button(fac, image = p_d, command=reject, bg = "white", highlightthickness = 0, bd = 0)
        bReject.place(x = 750, y = 550)

        bnext = Button(fac, image = p_n, command=next_file, bg = "white", highlightthickness = 0, bd = 0)
        bnext.place(x = 1150, y = 550)
        
        b_o = Button(fac, image = p_o, command=open_file, bg = "white", highlightthickness = 0, bd = 0)
        b_o.place(x = 1150, y = 250)

        global tCom
        # Creating a photoimage object to use image 
        tCom =  Text(fac, height=3, width=70)
        #eCom.insert(10, "COMMENTS")
        tCom.place(x = 50, y = 450)
        tCom.focus()

def department(*args):
    a = variable_dept.get()
    ld["text"] = a
    if (a=="EEE"): a= "C"   # options for f_eee
    elif (a=="ECE"): a= "D"
    elif (a=="CSE"): a= "E"
    elif (a=="ITD"): a= "F"
      
    opt_dept.place_forget()
    return a

    
#initialize the excel worksheets
sheet = wb['apple']
p_s = PhotoImage(file = r"C:\Users\Rohit\Desktop\Pro\start.png")
bs = Button(fac, image = p_s, command=start, bg = "white", highlightthickness = 0, bd = 0)
bs.place(x =650, y = 350) 

p_a = PhotoImage(file = r"C:\Users\Rohit\Desktop\Pro\accept_s.png")
p_n = PhotoImage(file = r"C:\Users\Rohit\Desktop\Pro\next.png")
p_d = PhotoImage(file = r"C:\Users\Rohit\Desktop\Pro\declined_s.png")
p_o = PhotoImage(file = r"C:\Users\Rohit\Desktop\Pro\open_f.png")
p_dd = PhotoImage(file = r"C:\Users\Rohit\Desktop\Pro\dd.png")

# select DEPARTMENT
lsel =Label(fac,text="SELECT DEPARTMENT", wraplength=600, background = "white", fg="black", justify="left",font = "times 18 bold")
lsel.place(x=320, y=205)
variable_dept = StringVar()
OptionList_dept = ["EEE", "ECE", "CSE", "ITD" ] 
variable_dept.set(OptionList_dept[0])
ld =Label(fac,text="", wraplength=600, background = "white", fg="black", justify="left",font = "times 18")
ld.place(x=680, y=205)
# options for DEPT
opt_dept = OptionMenu(fac, variable_dept, *OptionList_dept)
opt_dept.config(width=140, font=('Helvetica', 12))
opt_dept.place(x=680, y=200)
opt_dept.config(image =p_dd)
opt_dept.config(border = 0)
opt_dept["highlightthickness"]=0
opt_dept.config( bg = "white")
opt_dept["menu"].config(bg="white")

variable_dept.trace("w", department)
fac.mainloop()
