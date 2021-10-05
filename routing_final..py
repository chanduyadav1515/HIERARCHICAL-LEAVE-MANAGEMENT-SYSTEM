from tkinter import *
from openpyxl import load_workbook
from tkinter import messagebox
from datetime import date
from tkinter import filedialog
from docx import Document
import requests
import os

try:
    import tkinter as tk
    from tkinter import ttk
except ImportError:
    import Tkinter as tk
    import ttk
from datetime import timedelta, date
from tkcalendar import Calendar, DateEntry
import datetime 
import calendar 
from datetime import datetime
import pandas as pd
import smtplib 
from email.mime.multipart import MIMEMultipart 
from email.mime.text import MIMEText 
from email.mime.base import MIMEBase 
from email import encoders 

today_date = datetime.today().strftime('%Y-%m-%d')   # get today's date
print(today_date)
main = Tk()
main.geometry('1366x768')
main.configure(bg="white")
main.title("SUBMISSION OF LETTER")
fac = Tk()
fac.geometry('650x400')
fac.configure(bg="black")
fac.title("STATUS")
wb = load_workbook('index.xlsx')
sheet = wb['apple']
row_no = sheet.max_row+1  # to write data into that row

# select DEPARTMENT
variable_dept = StringVar()
OptionList_dept = ["EEE", "ECE", "CSE", "ITD" ] 
variable_dept.set(OptionList_dept[0])
# select FACULTY-EEE
variable_f = StringVar()
OptionList= ["SELECT"] 
variable_f.set(OptionList[0])
# select CATEGORY
variable_cat= StringVar()
OptionList_cat = ["CL ", "CCL", "EL ", "OD " ] 
variable_cat.set(OptionList_cat[0])
# select HOD
variable_hod = StringVar()
OptionList_hod = ["EEE", "ECE", "CSE", "ITD"] 
variable_hod.set(OptionList_hod[0])
# select DEAN
variable_dean = StringVar()
OptionList_dean = ["ACADEMICS", "STUDENT AFFAIRS", "RESEARCH", "FACULTY"] 
variable_dean.set(OptionList_dean[0])

# options for DEPT
opt_dept = OptionMenu(main, variable_dept, *OptionList_dept)
opt_dept.config(width=5, font=('Helvetica', 12))
opt_dept.place(x=220, y=650)
opt_dept.config(bg = "white") 
opt_dept["menu"].config(bg="white")
# options for HOD
opt_hod = OptionMenu(main, variable_hod, *OptionList_hod)
opt_hod.config(width=5, font=('Helvetica', 12))
opt_hod.place(x=760, y=500)
opt_hod.config(bg = "white") 
opt_hod["menu"].config(bg="white")
opt_hod['state'] = DISABLED
# options for DEAN
opt_dean = OptionMenu(main, variable_dean, *OptionList_dean)
opt_dean.config(width=15, font=('Helvetica', 12))
opt_dean.place(x=680, y=362)
opt_dean.config(bg = "white") 
opt_dean["menu"].config(bg="white")
opt_dean['state'] = DISABLED

# LABELS ----------------------------
label_F =Label(main, text="FACULTY DETAILS ", font=('Helvetica ', 14 ,'underline'), fg='orangered2', bg="white").place(x=605,y = 610)
label_dept =Label(main, text="Select Department", font=('Helvetica', 12), fg='blue',bg="white").place(x=20, y=655)
label_Hod =Label(main, text="HEAD OF THE DEPARTMENT  : ", font=('Helvetica ', 14,'bold'), fg='orangered3', bg="white").place(x=390, y=505)
#label_hod =Label(main, text="Select Department ", font=('Helvetica', 12), fg='blue', bg="white") .place(x=20, y=225)
label_Dean =Label(main, text="DEAN : ", font=('Helvetica ', 14,'bold'), fg='red2', bg="white").place(x=580, y=365)
label_Prin =Label(main, text="PRINCIPAL", font=('Helvetica ', 15 ,'bold'), fg='red3', bg="white").place(x=620, y=205)
label_Chrm =Label(main, text="CHAIRMAN", font=('Helvetica ', 16 ,'bold'), fg='red4', bg="white").place(x=620, y=75)
#label_hod =Label(main, text="Select Department ", font=('Helvetica', 12), fg='blue', bg="white").place(x=20, y=310)
label_temp =Label(main, text="", font=('Helvetica ', 14 ,'underline'), fg='white', bg="white")
label_temp.place(x=420, y=110)




    


subject = StringVar()
e1 = Entry(main,textvariable=subject,font=('Verdana',14),width=20).place(x = 250, y = 250)


from_label=Label(main, text='Choose from date',font=('Helvetica', 14),fg='blue',bg="white").place(x=150,y=100)
cal1 = DateEntry(main, width=12, background='darkblue',foreground='white', borderwidth=2)
cal1.place(x=350,y=100)
to_label=Label(main, text='Choose to date',font=('Helvetica', 14),fg='blue',bg="white").place(x=150,y=150)
cal2 = DateEntry(main, width=12, background='darkblue',foreground='white', borderwidth=2)
cal2.place(x=350,y=150)

sub_label=Label(main, text="Subject", font=('Helvetica', 16), fg='blue',bg="white").place(x=150,y=250)
#des_label=Label(main, text="Short Description", font=('Helvetica', 16), fg='blue',bg="white").place(x=150,y=300)

tCom =  Text(main, height=3, width=70)
tCom.pack()
#tCom.place(x = 350, y = 300)
sbt=Button(main, text='Submit', command=lambda: read_sub())
sbt.pack()
sbt.place(x=400,y=350)

submit=Button(main, text='Click Here', command=lambda: upload_mail())
submit.pack()
submit.place(x=600,y=450)




def department(*args):
    #labelTest.configure(text="The selected item is {}".format(variable1.get()))
    a = variable_dept.get()
    label_temp.config(text = label_temp['text']+ a +"_")
    label_faculty =Label(main, text="Select Faculty ID", font=('Helvetica', 12), fg='blue', bg="white")
    label_faculty .place(x=390, y=655)
    if (a=="EEE"): OptionList = ["01", "02", "03", "04" ]    # options for f_eee
    elif (a=="ECE"): OptionList = ["01", "02", "03", "04" ]    # options for f_ece
    elif (a=="CSE"): OptionList = ["01", "02", "03", "04" ]    # options for f_cse
    elif (a=="ITD"): OptionList = ["01", "02", "03", "04" ]    # options for f_IT
    global opt_f
    opt_f = OptionMenu(main, variable_f, *OptionList)
    opt_f.config(width=10, font=('Helvetica', 12))
    opt_f.place(x=570, y=655)
    opt_f.config(bg = "white") 
    opt_f["menu"].config(bg="white")
    opt_dept['state'] = DISABLED
    
def faculty(*args):
    a = variable_f.get()
    label_temp.config(text = label_temp['text']+ a +"_")
    label_cat =Label(main, text="Leave Type", font=('Helvetica', 12), fg='blue', bg="white")
    label_cat .place(x=920, y=650)
    global opt_cat
    opt_cat = OptionMenu(main, variable_cat, *OptionList_cat)
    opt_cat.config(width=5, font=('Helvetica', 12))
    opt_cat.place(x=1050, y=650)
    opt_cat.config(bg = "white") 
    opt_cat["menu"].config(bg="white")
    opt_f['state'] = DISABLED
    
def hod(*args):
    a = variable_hod.get()
    label_temp.config(text = label_temp['text']+ a +"_")
    opt_dean['state'] = NORMAL
    opt_hod['state'] = DISABLED 
    if (a=="EEE"): col = "C"+str(row_no)
    elif (a=="ECE"): col = "D"+str(row_no)
    elif (a=="CSE"):  col = "E"+str(row_no)
    elif (a=="ITD"):   col = "F"+str(row_no)
    print(col)
    sheet[col] = 1
def category(*args):
    a = variable_cat.get()
    label_temp.config(text = label_temp['text']+ a +"_")
    opt_hod['state'] = NORMAL
    opt_cat['state'] = DISABLED
   
def dean(*args):
    global string,s
    a = variable_dean.get()
    if (a=="ACADEMICS"):
        a = "ACA"
        col = "G"+str(row_no)
    elif (a=="STUDENT AFFAIRS"):
        a = "STA"
        col = "H"+str(row_no)
    elif (a=="RESEARCH"):
        a = "REA"
        col = "I"+str(row_no)
    elif (a=="FACULTY"):
        a = "FAC"
        col = "J"+str(row_no)
    label_temp.config(text = label_temp['text']+ a +"_")
    row = "A"+str(row_no)
    sheet[row] = label_temp['text'] +str(row_no)
    sheet[col] = 0
    row = "K"+str(row_no) # for principal
    sheet[row] = 0
    row = "L"+str(row_no) # for chairman
    sheet[row] = 0
    row = "b"+str(row_no) # for faculty
    sheet[row] = 9
    row = "M"+str(row_no) # fac sub date
    sheet[row] = str(today_date)
    
    #messagebox.showinfo("Submission"," Please Select the Letter to be submitted. " )
    wb.save('index.xlsx')
    opt_dean['state'] =  DISABLED        
    file1=open('index.xlsx','rb')
    server_url = 'https://tejusolutions.com/rohit/index.php'
    files = {'photo':file1}
    data = {'submit':True}
    response = requests.post(server_url, files=files, data=data)   
    
    #main.destroy()
def read_sub():
    global a,coment,todate,fromdate,string,s
    a=subject.get()
    print(a)
    coment = tCom.get("1.0","end-1c")
    print(coment)
    todate=cal1.get_date()
    fromdate=cal2.get_date()
    print(todate)
    print(fromdate)
    stringdate="TO DATE    :" +'\n'+str(todate)
    stringdate1="FROM DATE    :"+'\n' +str(fromdate)
    global document      
    document = Document()
    document.add_heading('Leave Application', 0)
    s='ID    :      '+label_temp['text']+str(row_no)
    string=label_temp['text'] +str(row_no)+".docx"
    k=label_temp['text']
    l=k[4:6]
    print("ID :"+l)
    document.add_paragraph(s)
    document.add_paragraph(str("FACULTY ID    :"))
    document.add_paragraph(l)
    document.add_paragraph(str("SUBMITTED DATE    :"))
    document.add_paragraph(str(today_date))
    document.add_paragraph(str("TO DATE    :" ))
    document.add_paragraph(str(todate))
    document.add_paragraph(str("FROM DATE    :" ))
    document.add_paragraph(str(fromdate))
    
    
    v=fromdate-todate
    nodays="NO. of Days    :"+'\n'+str(v.days)
    print(nodays)
    document.add_paragraph(str("NO. of Days    :"))
    document.add_paragraph(str(v.days))
    b='SUBJECT    :'
    document.add_paragraph(str(b))
    document.add_paragraph(str(a))
    document.add_heading('Content', 0)

    messagebox.showinfo("Content"," Please add content to your file. " )
    
    #document.add_paragraph(str(coment))
    document.save("C:\\Users\\Rohit\\Desktop\\Pro\\" + label_temp['text'] +str(row_no)+".docx")
    
    os.startfile(label_temp['text'] +str(row_no)+".docx")

    info()


def info():
    global nodays
    a=subject.get()
    print(a)
    print(coment)
    print(todate)
    stringdate=str(todate)
    a1=datetime.strptime(stringdate,'%Y-%m-%d').weekday()
    print(calendar.day_name[a1])
    
    string=datetime.today().strftime('%Y-%m-%d')
    list1=[]
    list1=stringdate.split("-")
    print(list1)
    start_dt = fromdate
    end_dt = todate
    '''for dt in daterange(start_dt, end_dt):
        print(dt.strftime("%Y-%m-%d"))'''

    #print(string)
    v=fromdate-todate
    nodays="NO. of Days    :"+str(v.days)
    print(nodays)
    document.add_paragraph(str(nodays))
    
    lDate=Label(fac, text ="Date Submitted  : ", background = "white", fg="blue",font = "Helvetica 12")
    lDate.place(x=150, y=20)
    lDate_text=Label(fac,text=string, background = "white", fg="black", justify="left",font = "times 12")
    lDate_text.place(x=300, y=20)
    lDate=Label(fac, text ="From DATE       : ", background = "white", fg="blue",font = "Helvetica 12")
    lDate.place(x=150, y=70)
    lDate_text=Label(fac,text=todate, background = "white", fg="black", justify="left",font = "times 12")
    lDate_text.place(x=300, y=70)
    lDate=Label(fac, text ="To DATE         : ", background = "white", fg="blue",font = "Helvetica 12")
    lDate.place(x=150, y=120)
    lDate_text=Label(fac,text=fromdate, background = "white", fg="black", justify="left",font = "times 12")
    lDate_text.place(x=300, y=120)

    lDate=Label(fac, text ="No. of Days        : ", background = "white", fg="blue",font = "Helvetica 12")
    lDate.place(x=150, y=170)
    lDate_text=Label(fac,text=nodays, background = "white", fg="black", justify="left",font = "times 12")
    lDate_text.place(x=300, y=170)

    '''lSub1=Label(fac, text ="CONTENT           : ", background = "white", fg="green",font = "Helvetica 14 bold ")
    lSub1.place(x=150, y=250)
    lSubject1=Label(fac,text=coment, wraplength=1000, background = "white", fg="OrangeRed3", justify="left",font = "times 14")
    lSubject1.place(x=400, y=250)'''
                
    lSub=Label(fac, text ="SUBJECT            : ", background = "white", fg="firebrick3",font = "Helvetica 14 bold ")
    lSub.place(x=150, y=200)
    lSubject=Label(fac,text=a, wraplength=1000, background = "white", fg="OrangeRed3", justify="left",font = "times 14")
    lSubject.place(x=350, y=200)
    
def upload_mail():
    #document.save("C:\\Users\\Rohit\\Desktop\\Pro\\" + label_temp['text'] +str(row_no)+".docx")
    
    
    fromaddr = "rohitpvs2000@gmail.com"
    toaddr = "phani5016@gmail.com"
    msg = MIMEMultipart() 
    msg['From'] = fromaddr 
    msg['To'] = toaddr 
    msg['Subject'] = "Your application has been forwarded successfully." 
    body = "thank you" 
    msg.attach(MIMEText(body, 'plain')) 
    # open the file to be sent  
    filename = label_temp['text'] +str(row_no)+".docx"
    #attachment = open("C:\\Users\\Rohit\\Desktop\\", "rb")
    attachment=open(string,'rb')
    p = MIMEBase('application', 'octet-stream') 
    p.set_payload((attachment).read()) 
    encoders.encode_base64(p) 
    p.add_header('Content-Disposition', "attachment; filename= %s" % filename) 
    msg.attach(p) 
    s = smtplib.SMTP('smtp.gmail.com', 587) 
    s.starttls()  
    s.login(fromaddr, "rohit123") 
    text = msg.as_string() 
    s.sendmail(fromaddr, toaddr, text) 
    s.quit()
    messagebox.showinfo("Please Wait"," Your application is being uploaded. " )
    file2=open(string,'rb')
    server_url = 'https://tejusolutions.com/rohit/index.php'
    files = {'photo':file2}
    data = {'submit':True}
    response = requests.post(server_url, files=files, data=data)
    messagebox.showinfo("Succesful"," Your application has been forwarded successfully. " )
    main.destroy()
    



variable_dept.trace("w", department)
variable_f.trace("w", faculty)
variable_hod.trace("w", hod)
variable_cat.trace("w", category)
variable_dean.trace("w", dean)

main.mainloop()
