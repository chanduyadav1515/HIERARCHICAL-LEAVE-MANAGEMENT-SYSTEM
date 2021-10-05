from docx import Document
import os
global document      
document = Document()
s='hi'+'\n'+'bye'
document.add_heading('Leave Application', 0)
document.add_paragraph(str(s))
document.add_paragraph(str('hello'))
document.add_paragraph(str('done'))  
document.add_paragraph(str(7))
document.add_heading('Content', 0)  
document.add_paragraph(str('coment'))
document.save("C:\\Users\\Rohit\\Desktop\\Pro\\" + "temp.docx")

document = Document('temp.docx')
os.system('start temp.docx')

sub = document.paragraphs[1].text
body = document.paragraphs[2].text
date = document.paragraphs[3].text
From = document.paragraphs[4].text
x=document.paragraphs[6].text
document.save("C:\\Users\\Rohit\\Desktop\\Pro\\" + "temp.docx")
print(sub)
print(body)
print(date)
print(From)
print(x)


